import os
from dotenv import load_dotenv
from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, filters, ContextTypes
from langchain.document_loaders import TextLoader
import pdfplumber
from langchain.schema import Document
from langchain.embeddings import OpenAIEmbeddings
from langchain.vectorstores import Chroma
from langchain.chat_models import ChatOpenAI
from langchain.chains import RetrievalQA
import docx
from langchain.prompts import PromptTemplate

load_dotenv()
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
CHROMA_DIR = "chroma_index"
DOCS_DIR = "docs"

# Prompt customizado para garantir que só use os documentos indexados
def get_custom_prompt():
    return PromptTemplate(
        input_variables=["context", "question"],
        template=(
            "Responda à pergunta apenas com base nos documentos fornecidos abaixo. "
            "Se a resposta não estiver nos documentos, diga 'Não sei com base nos documentos indexados.'\n\n"
            "Documentos:\n{context}\n\nPergunta: {question}\nResposta:"
        )
    )

# Função para carregar PDF usando pdfplumber
def load_pdf_with_pdfplumber(path):
    docs = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text()
            if text:
                docs.append(Document(page_content=text, metadata={"source": path, "page": i+1}))
    return docs

# Função para carregar DOCX usando python-docx
def load_docx_with_python_docx(path):
    doc = docx.Document(path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
    if full_text:
        return [Document(page_content="\n".join(full_text), metadata={"source": path})]
    else:
        return []

# load or create Chroma index
def get_retriever():
    embeddings = OpenAIEmbeddings()
    if os.path.exists(CHROMA_DIR) and os.listdir(CHROMA_DIR):
        vectorstore = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)
    else:
        vectorstore = Chroma.from_documents([], embeddings, persist_directory=CHROMA_DIR)
    return vectorstore.as_retriever()

def add_docs_to_index(docs):
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma(persist_directory=CHROMA_DIR, embedding_function=embeddings)
    vectorstore.add_documents(docs)
    vectorstore.persist()

retriever = get_retriever()
custom_prompt = get_custom_prompt()
qa = RetrievalQA.from_chain_type(
    llm=ChatOpenAI(model="gpt-4.1-nano"),
    retriever=retriever,
    chain_type="stuff",
    chain_type_kwargs={"prompt": custom_prompt}
)

async def start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text("Olá! Envie um arquivo .txt, .pdf ou .docx para indexar, ou faça uma pergunta.")

async def handle_docs(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    file = update.message.document
    ext = os.path.splitext(file.file_name)[1].lower()
    if ext not in [".txt", ".pdf", ".docx"]:
        await update.message.reply_text("Arquivo não suportado. Envie apenas .txt, .pdf ou .docx.")
        return
    os.makedirs(DOCS_DIR, exist_ok=True)
    path = f"{DOCS_DIR}/{file.file_name}"
    try:
        telegram_file = await file.get_file()
        await telegram_file.download_to_drive(path)
        if ext == ".txt":
            loader = TextLoader(path)
            docs = loader.load()
        elif ext == ".pdf":
            docs = load_pdf_with_pdfplumber(path)
        elif ext == ".docx":
            docs = load_docx_with_python_docx(path)
        else:
            docs = []
        add_docs_to_index(docs)
        global retriever, qa
        retriever = get_retriever()
        qa = RetrievalQA.from_chain_type(
            llm=ChatOpenAI(model="gpt-4.1-nano"),
            retriever=retriever,
            chain_type="stuff",
            chain_type_kwargs={"prompt": custom_prompt}
        )
        await update.message.reply_text(f"Arquivo indexado com sucesso: {file.file_name} (total de {len(docs)} documentos)")
    except Exception as e:
        await update.message.reply_text(f"Erro ao processar o arquivo: {str(e)}")

async def answer(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    question = update.message.text
    try:
        response = qa.run(question)
        await update.message.reply_text(response)
    except Exception as e:
        await update.message.reply_text(f"Erro ao buscar resposta: {str(e)}")

app = ApplicationBuilder().token(TELEGRAM_TOKEN).build()
app.add_handler(CommandHandler("start", start))
app.add_handler(MessageHandler(filters.Document.ALL, handle_docs))
app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, answer))
app.run_polling()
